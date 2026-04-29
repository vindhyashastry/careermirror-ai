import os
import re
import fitz  # PyMuPDF
import docx
import spacy
from .ai_utils import get_embedding_model

# ──────────────────────── Configuration ────────────────────────
try:
    nlp = spacy.load("en_core_web_md")
except:
    # Fallback if model not downloaded
    print("⏳ Downloading SpaCy model (First time only)...")
    os.system("python -m spacy download en_core_web_md")
    nlp = spacy.load("en_core_web_md")

# ──────────────────────── Skill Taxonomy ────────────────────────
SKILL_TAXONOMY = {
    "languages": [
        "Python", "JavaScript", "Java", "C++", "C#", "TypeScript", "Go", "Rust", "Ruby", "PHP", 
        "Swift", "Kotlin", "Scala", "C", "SQL", "R", "Dart", "HTML", "CSS", "Bash", "Shell"
    ],
    "frameworks": [
        "React", "Node.js", "Django", "Flask", "FastAPI", "Spring Boot", "Angular", "Vue.js", "Next.js", 
        "Express", "Flutter", "React Native", "TensorFlow", "PyTorch", "Laravel", "Svelte", "Redux"
    ],
    "tools": [
        "AWS", "Docker", "Kubernetes", "Git", "GitHub", "Jenkins", "Terraform", "Linux", 
        "PostgreSQL", "MySQL", "MongoDB", "Redis", "Elasticsearch", "Firebase", "Azure", "GCP",
        "JWT", "OAuth2", "REST API", "GraphQL", "Microservices", "Docker Compose", "Nginx", "RabbitMQ"
    ]
}

# Create a flattened lower-case lookup map for speed
ALL_SKILLS_LOWER = {}
for category, skills in SKILL_TAXONOMY.items():
    for skill in skills:
        ALL_SKILLS_LOWER[skill.lower()] = (skill, category)

# ──────────────────────── Section Patterns ────────────────────────
SECTION_PATTERNS = {
    "contact": r"contact|details|information",
    "education": r"education|academic|studies",
    "experience": r"experience|employment|work|history",
    "skills": r"skills|competencies|technologies|stack",
    "projects": r"projects|work samples",
    "certifications": r"certifications|licenses|awards|courses",
    "summary": r"summary|objective|profile",
}

# ──────────────────────── Core Parser Class ────────────────────────
class ResumeParser:

    def __init__(self):
        self.nlp = nlp

    # ──────────────── Text Extraction ────────────────
    def extract_text(self, file_path: str) -> str:
        """Extract raw text from PDF or DOCX/DOC."""
        lower = file_path.lower()
        if lower.endswith(".pdf"):
            return self._extract_pdf(file_path)
        elif lower.endswith(".docx") or lower.endswith(".doc"):
            return self._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_path}")

    def _extract_pdf(self, path: str) -> str:
        text = ""
        with fitz.open(path) as doc:
            for page in doc:
                text += page.get_text() + "\n"
        return text.strip()

    def _extract_docx(self, path: str) -> str:
        doc = docx.Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also grab table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)

    # ──────────────── Section Detection ────────────────
    def detect_sections(self, text: str) -> dict:
        """Split the resume text into labelled sections."""
        lines = text.split("\n")
        sections = {k: [] for k in SECTION_PATTERNS}
        sections["other"] = []
        current = "other"

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            matched = False
            for section, pattern in SECTION_PATTERNS.items():
                if re.search(pattern, stripped, re.IGNORECASE) and len(stripped) < 60:
                    current = section
                    matched = True
                    break
            if not matched:
                sections[current].append(stripped)

        return {k: "\n".join(v) for k, v in sections.items() if v}

    # ──────────────── Skill Extraction ────────────────
    def extract_skills(self, text: str) -> dict:
        """
        Extract and categorize all skills found in the text.
        Returns {'languages': [...], 'frameworks': [...], 'tools': [...], 'all': [...]}
        """
        found: dict = {"languages": set(), "frameworks": set(), "tools": set()}
        text_lower = text.lower()

        for skill_lower, (canonical, category) in ALL_SKILLS_LOWER.items():
            # Word-boundary match to avoid partial hits (e.g. 'C' inside 'CI/CD')
            pattern = r"(?<![a-zA-Z0-9\-+#])" + re.escape(skill_lower) + r"(?![a-zA-Z0-9\-+#])"
            if re.search(pattern, text_lower):
                found[category].add(canonical)

        return {
            "languages":  sorted(found["languages"]),
            "frameworks": sorted(found["frameworks"]),
            "tools":      sorted(found["tools"]),
            "all":        sorted(found["languages"] | found["frameworks"] | found["tools"]),
        }

    # ──────────────── Contact Info ────────────────
    def extract_contact(self, text: str) -> dict:
        emails = re.findall(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", text)
        phones = re.findall(r"(\+?\d[\d\s\-().]{7,}\d)", text)
        linkedin = re.findall(r"linkedin\.com/in/[\w\-]+", text, re.IGNORECASE)
        github = re.findall(r"github\.com/[\w\-]+", text, re.IGNORECASE)
        return {
            "emails": list(set(emails)),
            "phones": [p.strip() for p in phones[:3]],
            "linkedin": linkedin[0] if linkedin else None,
            "github": github[0] if github else None,
        }

    # ──────────────── Section-Specific Extractors ────────────────
    def extract_education(self, sections: dict) -> list:
        edu_text = sections.get("education", "")
        degrees = re.findall(
            r"(B\.?S\.?|M\.?S\.?|B\.?E\.?|M\.?E\.?|B\.?Tech|M\.?Tech|PhD|B\.?Sc|M\.?Sc|Bachelor|Master|Doctorate)[^,\n]*",
            edu_text, re.IGNORECASE
        )
        years = re.findall(r"\b(19|20)\d{2}\b", edu_text)
        institutions = re.findall(r"(?:University|College|Institute|School|IIT|NIT|MIT|IIM)[^\n,]*", edu_text, re.IGNORECASE)
        return [{"degree": d.strip(), "institutions": institutions, "years": years} for d in degrees] if degrees else [{"raw": edu_text}]

    def extract_experience(self, sections: dict) -> list:
        exp_text = sections.get("experience", "")
        entries = []
        # Split by year ranges which usually mark job boundaries
        parts = re.split(r"\n(?=\d{4}|\w+ \d{4}|Present)", exp_text)
        for part in parts:
            if part.strip():
                years = re.findall(r"\b(19|20)\d{2}\b", part)
                entries.append({
                    "text": part.strip()[:300],
                    "years": list(set(years)),
                })
        return entries if entries else [{"raw": exp_text[:500]}]

    def extract_projects(self, sections: dict) -> list:
        proj_text = sections.get("projects", "")
        projects = []
        # Split on lines that start with a capital word (project name pattern)
        entries = re.split(r"\n(?=[A-Z][^\n]{3,60}\n)", proj_text)
        for entry in entries:
            if entry.strip():
                lines = [l.strip() for l in entry.split("\n") if l.strip()]
                name = lines[0] if lines else "Project"
                description = " ".join(lines[1:])[:400] if len(lines) > 1 else ""
                # Extract tech used in project
                tech = self.extract_skills(entry)["all"]
                projects.append({"name": name, "description": description, "technologies": tech})
        return projects if projects else []

    def extract_certifications(self, sections: dict) -> list:
        cert_text = sections.get("certifications", "")
        certs = []
        for line in cert_text.split("\n"):
            line = line.strip()
            if line and len(line) > 5:
                year = re.search(r"\b(19|20)\d{2}\b", line)
                certs.append({"name": line, "year": year.group() if year else None})
        return certs

    # ──────────────── Keyword Normalization ────────────────
    def normalize_keywords(self, skills: list) -> list:
        """Lowercase, strip, deduplicate, then re-map to canonical name."""
        seen = set()
        normalized = []
        for skill in skills:
            key = skill.strip().lower()
            if key and key not in seen:
                seen.add(key)
                # Use canonical form if available
                canonical = ALL_SKILLS_LOWER.get(key, (skill.strip(), None))[0]
                normalized.append(canonical)
        return normalized

    # ──────────────── Embedding Generation ────────────────
    def generate_embedding(self, text: str) -> list:
        """Generate a 384-dim semantic embedding for similarity matching."""
        # Use first 512 words to keep it fast
        model = get_embedding_model()
        if not model:
            return [0.0] * 384
            
        truncated = " ".join(text.split()[:512])
        embedding = model.encode(truncated, normalize_embeddings=True)
        return embedding.tolist()

    # ──────────────── Master Parse ────────────────
    def parse(self, file_path: str) -> tuple:
        """
        Full pipeline. Returns (raw_text, structured_data_dict).
        structured_data_dict contains all extracted fields + embedding.
        """
        # 1. Extract raw text
        raw_text = self.extract_text(file_path)

        # 2. Detect sections
        sections = self.detect_sections(raw_text)

        # 3. Extract all entities
        skills_data = self.extract_skills(raw_text)
        contact = self.extract_contact(raw_text)
        education = self.extract_education(sections)
        experience = self.extract_experience(sections)
        projects = self.extract_projects(sections)
        certifications = self.extract_certifications(sections)

        # 4. Normalize & deduplicate skills
        normalized_skills = self.normalize_keywords(skills_data["all"])

        # 5. Generate embedding
        embedding = self.generate_embedding(raw_text)

        structured = {
            "contact": contact,
            "skills": normalized_skills,  # flat deduplicated list
            "skill_categories": {
                "languages": skills_data["languages"],
                "frameworks": skills_data["frameworks"],
                "tools": skills_data["tools"],
            },
            "education": education,
            "experience": experience,
            "projects": projects,
            "certifications": certifications,
            "sections_detected": list(sections.keys()),
            "embedding": embedding,         # 384-dim float list
            "raw_text_preview": raw_text[:800] + ("..." if len(raw_text) > 800 else ""),
        }

        return raw_text, structured
